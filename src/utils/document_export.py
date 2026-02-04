"""
Document export utilities for AI Practice Platform.

Provides export functionality for PDF, DOCX, and other formats.
"""

import os
import re
from io import BytesIO
from datetime import datetime
from typing import Optional, Dict, Any, List

# PDF generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    PageBreak, ListFlowable, ListItem
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

# DOCX generation
from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


class PDFExporter:
    """Export documents to PDF format."""

    def __init__(self):
        self.styles = getSampleStyleSheet()
        self._setup_custom_styles()

    def _setup_custom_styles(self):
        """Set up custom PDF styles."""
        # Title style
        self.styles.add(ParagraphStyle(
            name='CustomTitle',
            parent=self.styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.HexColor('#1e3a5f'),
            alignment=TA_CENTER
        ))

        # Heading 1
        self.styles.add(ParagraphStyle(
            name='CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceBefore=20,
            spaceAfter=12,
            textColor=colors.HexColor('#2c5282')
        ))

        # Heading 2
        self.styles.add(ParagraphStyle(
            name='CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
            textColor=colors.HexColor('#4299e1')
        ))

        # Body text
        self.styles.add(ParagraphStyle(
            name='CustomBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            leading=14
        ))

        # Bullet style
        self.styles.add(ParagraphStyle(
            name='CustomBullet',
            parent=self.styles['Normal'],
            fontSize=11,
            leftIndent=20,
            spaceAfter=4
        ))

    def export(
        self,
        title: str,
        content: str,
        organization: str = "Organization",
        metadata: Optional[Dict] = None
    ) -> BytesIO:
        """
        Export content to PDF.

        Args:
            title: Document title
            content: Markdown content to convert
            organization: Organization name for header
            metadata: Optional metadata dict

        Returns:
            BytesIO buffer containing the PDF
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        story = []

        # Title
        story.append(Paragraph(title, self.styles['CustomTitle']))
        story.append(Spacer(1, 12))

        # Organization and date
        date_str = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(
            f"Prepared for: {organization}",
            self.styles['Normal']
        ))
        story.append(Paragraph(f"Date: {date_str}", self.styles['Normal']))
        story.append(Spacer(1, 24))

        # Parse and add content
        self._add_markdown_content(story, content)

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer

    def _add_markdown_content(self, story: List, content: str):
        """Parse markdown content and add to PDF story."""
        lines = content.split('\n')
        in_list = False
        list_items = []

        for line in lines:
            line = line.strip()

            if not line:
                if in_list and list_items:
                    story.append(ListFlowable(
                        [ListItem(Paragraph(item, self.styles['CustomBullet']))
                         for item in list_items],
                        bulletType='bullet'
                    ))
                    list_items = []
                    in_list = False
                story.append(Spacer(1, 6))
                continue

            # Headers
            if line.startswith('# '):
                story.append(Paragraph(
                    self._clean_markdown(line[2:]),
                    self.styles['CustomHeading1']
                ))
            elif line.startswith('## '):
                story.append(Paragraph(
                    self._clean_markdown(line[3:]),
                    self.styles['CustomHeading1']
                ))
            elif line.startswith('### '):
                story.append(Paragraph(
                    self._clean_markdown(line[4:]),
                    self.styles['CustomHeading2']
                ))
            elif line.startswith('#### '):
                story.append(Paragraph(
                    self._clean_markdown(line[5:]),
                    self.styles['CustomHeading2']
                ))
            # Lists
            elif line.startswith('- ') or line.startswith('* '):
                in_list = True
                list_items.append(self._clean_markdown(line[2:]))
            elif re.match(r'^\d+\.', line):
                in_list = True
                list_items.append(self._clean_markdown(re.sub(r'^\d+\.\s*', '', line)))
            else:
                # Regular paragraph
                story.append(Paragraph(
                    self._clean_markdown(line),
                    self.styles['CustomBody']
                ))

        # Flush remaining list items
        if list_items:
            story.append(ListFlowable(
                [ListItem(Paragraph(item, self.styles['CustomBullet']))
                 for item in list_items],
                bulletType='bullet'
            ))

    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting for PDF."""
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
        # Code
        text = re.sub(r'`(.+?)`', r'<font face="Courier">\1</font>', text)
        # Links - just show text
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text


class DOCXExporter:
    """Export documents to DOCX format."""

    def __init__(self):
        self.doc = None

    def export(
        self,
        title: str,
        content: str,
        organization: str = "Organization",
        metadata: Optional[Dict] = None
    ) -> BytesIO:
        """
        Export content to DOCX.

        Args:
            title: Document title
            content: Markdown content to convert
            organization: Organization name for header
            metadata: Optional metadata dict

        Returns:
            BytesIO buffer containing the DOCX
        """
        self.doc = DocxDocument()
        self._setup_styles()

        # Title
        title_para = self.doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Organization and date
        date_str = datetime.now().strftime("%B %d, %Y")
        meta_para = self.doc.add_paragraph()
        meta_para.add_run(f"Prepared for: {organization}").bold = True
        meta_para.add_run(f"\nDate: {date_str}")
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph()  # Spacing

        # Parse and add content
        self._add_markdown_content(content)

        # Save to buffer
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer

    def _setup_styles(self):
        """Configure document styles."""
        # Set default font
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def _add_markdown_content(self, content: str):
        """Parse markdown content and add to DOCX."""
        lines = content.split('\n')
        current_list = None

        for line in lines:
            original_line = line
            line = line.strip()

            if not line:
                if current_list:
                    current_list = None
                continue

            # Headers
            if line.startswith('# '):
                self.doc.add_heading(self._clean_markdown(line[2:]), 1)
            elif line.startswith('## '):
                self.doc.add_heading(self._clean_markdown(line[3:]), 1)
            elif line.startswith('### '):
                self.doc.add_heading(self._clean_markdown(line[4:]), 2)
            elif line.startswith('#### '):
                self.doc.add_heading(self._clean_markdown(line[5:]), 3)
            # Bullet lists
            elif line.startswith('- ') or line.startswith('* '):
                para = self.doc.add_paragraph(
                    self._clean_markdown(line[2:]),
                    style='List Bullet'
                )
            # Numbered lists
            elif re.match(r'^\d+\.', line):
                text = re.sub(r'^\d+\.\s*', '', line)
                para = self.doc.add_paragraph(
                    self._clean_markdown(text),
                    style='List Number'
                )
            else:
                # Regular paragraph
                para = self.doc.add_paragraph()
                self._add_formatted_text(para, line)

    def _add_formatted_text(self, paragraph, text: str):
        """Add text with markdown formatting to paragraph."""
        # Simple formatting parser
        parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`.*?`)', text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = paragraph.add_run(part[1:-1])
                run.font.name = 'Courier New'
            else:
                # Remove link markdown, keep text
                clean = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', part)
                paragraph.add_run(clean)

    def _clean_markdown(self, text: str) -> str:
        """Remove markdown formatting for plain text."""
        # Bold
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # Italic
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        # Code
        text = re.sub(r'`(.+?)`', r'\1', text)
        # Links - just show text
        text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
        return text


class DocumentExporter:
    """Main document exporter with format selection."""

    def __init__(self):
        self.pdf_exporter = PDFExporter()
        self.docx_exporter = DOCXExporter()

    def export(
        self,
        title: str,
        content: str,
        format: str = 'pdf',
        organization: str = "Organization",
        metadata: Optional[Dict] = None
    ) -> BytesIO:
        """
        Export document to specified format.

        Args:
            title: Document title
            content: Markdown content
            format: Export format ('pdf', 'docx', 'markdown')
            organization: Organization name
            metadata: Optional metadata

        Returns:
            BytesIO buffer with exported content
        """
        if format == 'pdf':
            return self.pdf_exporter.export(title, content, organization, metadata)
        elif format == 'docx':
            return self.docx_exporter.export(title, content, organization, metadata)
        elif format == 'markdown' or format == 'md':
            buffer = BytesIO()
            buffer.write(content.encode('utf-8'))
            buffer.seek(0)
            return buffer
        else:
            raise ValueError(f"Unsupported format: {format}")

    def get_mime_type(self, format: str) -> str:
        """Get MIME type for format."""
        mime_types = {
            'pdf': 'application/pdf',
            'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'markdown': 'text/markdown',
            'md': 'text/markdown'
        }
        return mime_types.get(format, 'application/octet-stream')

    def get_extension(self, format: str) -> str:
        """Get file extension for format."""
        extensions = {
            'pdf': '.pdf',
            'docx': '.docx',
            'markdown': '.md',
            'md': '.md'
        }
        return extensions.get(format, '.txt')


class AssessmentPDFExporter:
    """Export assessment reports to branded PDF format."""

    def __init__(self):
        self.styles = getSampleStyleSheet()

    def _hex_to_color(self, hex_color: str):
        """Convert hex color string to reportlab color."""
        hex_color = hex_color.lstrip('#')
        return colors.HexColor(f'#{hex_color}')

    def _setup_branded_styles(self, primary_color: str, secondary_color: str):
        """Set up custom PDF styles with branding colors."""
        primary = self._hex_to_color(primary_color)
        secondary = self._hex_to_color(secondary_color)

        # Create new styles each time to support different branding
        custom_styles = {}

        custom_styles['BrandedTitle'] = ParagraphStyle(
            name='BrandedTitle',
            parent=self.styles['Title'],
            fontSize=28,
            spaceAfter=20,
            textColor=primary,
            alignment=TA_CENTER
        )

        custom_styles['BrandedSubtitle'] = ParagraphStyle(
            name='BrandedSubtitle',
            parent=self.styles['Normal'],
            fontSize=14,
            spaceAfter=30,
            textColor=colors.grey,
            alignment=TA_CENTER
        )

        custom_styles['BrandedHeading1'] = ParagraphStyle(
            name='BrandedHeading1',
            parent=self.styles['Heading1'],
            fontSize=18,
            spaceBefore=24,
            spaceAfter=12,
            textColor=primary
        )

        custom_styles['BrandedHeading2'] = ParagraphStyle(
            name='BrandedHeading2',
            parent=self.styles['Heading2'],
            fontSize=14,
            spaceBefore=16,
            spaceAfter=8,
            textColor=secondary
        )

        custom_styles['BrandedBody'] = ParagraphStyle(
            name='BrandedBody',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=8,
            alignment=TA_JUSTIFY,
            leading=14
        )

        custom_styles['BrandedScore'] = ParagraphStyle(
            name='BrandedScore',
            parent=self.styles['Normal'],
            fontSize=48,
            textColor=primary,
            alignment=TA_CENTER
        )

        custom_styles['BrandedLabel'] = ParagraphStyle(
            name='BrandedLabel',
            parent=self.styles['Normal'],
            fontSize=12,
            textColor=colors.grey,
            alignment=TA_CENTER
        )

        return custom_styles

    def export(
        self,
        assessment_data: Dict[str, Any],
        branding: Dict[str, Any],
        organization_name: str = "Organization"
    ) -> BytesIO:
        """
        Export assessment to branded PDF.

        Args:
            assessment_data: Assessment result data
            branding: Organization branding settings (colors, logo, name)
            organization_name: Organization name

        Returns:
            BytesIO buffer containing the PDF
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=60,
            leftMargin=60,
            topMargin=60,
            bottomMargin=60
        )

        primary_color = branding.get('primary_color', '#0066FF')
        secondary_color = branding.get('secondary_color', '#00D4AA')
        brand_name = branding.get('name', 'AI Practice Platform')
        powered_by = branding.get('powered_by', 'Patriot Tech Systems Consulting LLC')

        styles = self._setup_branded_styles(primary_color, secondary_color)
        story = []

        # Header with branding
        story.append(Paragraph(brand_name, styles['BrandedTitle']))
        story.append(Paragraph("AI Readiness Assessment Report", styles['BrandedSubtitle']))
        story.append(Spacer(1, 10))

        # Organization and date
        date_str = datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(
            f"<b>Prepared for:</b> {organization_name}",
            styles['BrandedBody']
        ))
        story.append(Paragraph(f"<b>Date:</b> {date_str}", styles['BrandedBody']))
        story.append(Spacer(1, 30))

        # Extract result data
        result = assessment_data.get('result', assessment_data)
        overall_score = result.get('overall_score', 0)
        maturity_level = result.get('maturity_level', 'Exploring')

        # Overall Score Section
        story.append(Paragraph("Executive Summary", styles['BrandedHeading1']))
        story.append(Spacer(1, 10))

        # Score display in a table for better formatting
        score_data = [
            [Paragraph(f"{int(overall_score)}", styles['BrandedScore'])],
            [Paragraph("Overall AI Readiness Score", styles['BrandedLabel'])],
            [Paragraph(f"Maturity Level: <b>{maturity_level}</b>", styles['BrandedBody'])]
        ]
        score_table = Table(score_data, colWidths=[300])
        score_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(score_table)
        story.append(Spacer(1, 20))

        # Dimension Scores
        dimension_scores = result.get('dimension_scores', {})
        if dimension_scores:
            story.append(Paragraph("Assessment Dimensions", styles['BrandedHeading1']))
            story.append(Spacer(1, 10))

            dim_data = [['Dimension', 'Score', 'Status']]
            for dim_id, dim_info in dimension_scores.items():
                dim_name = dim_info.get('name', dim_id.replace('_', ' ').title())
                dim_score = dim_info.get('score', 0)
                if dim_score >= 70:
                    status = 'Strong'
                elif dim_score >= 40:
                    status = 'Developing'
                else:
                    status = 'Needs Attention'
                dim_data.append([dim_name, f"{int(dim_score)}%", status])

            dim_table = Table(dim_data, colWidths=[250, 80, 100])
            dim_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), self._hex_to_color(primary_color)),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('ALIGN', (1, 0), (2, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('FONTSIZE', (0, 1), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 0), (-1, 0), 12),
                ('TOPPADDING', (0, 1), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 1), (-1, -1), 8),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8f9fa')]),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dee2e6')),
            ]))
            story.append(dim_table)
            story.append(Spacer(1, 20))

        # Strengths
        strengths = result.get('strengths', [])
        if strengths:
            story.append(Paragraph("Key Strengths", styles['BrandedHeading1']))
            story.append(Spacer(1, 10))
            for s in strengths[:5]:
                area = s.get('area', '')
                desc = s.get('description', '')
                story.append(Paragraph(f"<b>✓ {area}</b>", styles['BrandedBody']))
                story.append(Paragraph(f"   {desc}", styles['BrandedBody']))
            story.append(Spacer(1, 15))

        # Improvements
        improvements = result.get('improvements', [])
        if improvements:
            story.append(Paragraph("Areas for Improvement", styles['BrandedHeading1']))
            story.append(Spacer(1, 10))
            for i in improvements[:5]:
                area = i.get('area', '')
                rec = i.get('recommendation', '')
                story.append(Paragraph(f"<b>→ {area}</b>", styles['BrandedBody']))
                story.append(Paragraph(f"   {rec}", styles['BrandedBody']))
            story.append(Spacer(1, 15))

        # Recommendations
        recommendations = result.get('recommendations', [])
        if recommendations:
            story.append(PageBreak())
            story.append(Paragraph("Strategic Recommendations", styles['BrandedHeading1']))
            story.append(Spacer(1, 10))
            for idx, rec in enumerate(recommendations[:8], 1):
                if isinstance(rec, dict):
                    rec_text = rec.get('text', rec.get('recommendation', str(rec)))
                else:
                    rec_text = str(rec)
                story.append(Paragraph(f"{idx}. {rec_text}", styles['BrandedBody']))
            story.append(Spacer(1, 20))

        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph(
            f"<i>Report generated by {brand_name}</i>",
            ParagraphStyle('Footer', parent=self.styles['Normal'],
                          fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
        ))
        story.append(Paragraph(
            f"<i>Powered by {powered_by}</i>",
            ParagraphStyle('Footer2', parent=self.styles['Normal'],
                          fontSize=9, textColor=colors.grey, alignment=TA_CENTER)
        ))

        # Build PDF
        doc.build(story)
        buffer.seek(0)
        return buffer


# Singleton instances
_exporter = None
_assessment_exporter = None


def get_document_exporter() -> DocumentExporter:
    """Get singleton DocumentExporter instance."""
    global _exporter
    if _exporter is None:
        _exporter = DocumentExporter()
    return _exporter


def get_assessment_exporter() -> AssessmentPDFExporter:
    """Get singleton AssessmentPDFExporter instance."""
    global _assessment_exporter
    if _assessment_exporter is None:
        _assessment_exporter = AssessmentPDFExporter()
    return _assessment_exporter
